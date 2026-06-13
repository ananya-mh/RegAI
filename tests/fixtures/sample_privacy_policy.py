"""Generate a dummy company privacy policy DOCX for testing gap analysis."""

from pathlib import Path
from docx import Document
from docx.shared import Pt


def create_sample_policy(output_path: str = "sample_privacy_policy.docx") -> Path:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    doc.add_heading("Acme Corp — Data Privacy Policy", level=0)
    doc.add_paragraph("Version 2.3 | Effective Date: January 15, 2025")
    doc.add_paragraph("")

    doc.add_heading("1. Purpose and Scope", level=1)
    doc.add_paragraph(
        "This policy defines how Acme Corp collects, processes, stores, and shares "
        "personal data of customers, employees, and partners. It applies to all departments, "
        "contractors, and third-party processors acting on behalf of Acme Corp."
    )

    doc.add_heading("2. Data Collection", level=1)
    doc.add_paragraph(
        "Acme Corp collects the following categories of personal data: name, email address, "
        "phone number, billing address, IP address, and browser cookies. Data is collected "
        "through our website, mobile applications, customer support interactions, and "
        "third-party integrations."
    )
    doc.add_paragraph(
        "We collect data for the purposes of service delivery, billing, customer support, "
        "and product improvement. Marketing communications require explicit opt-in consent."
    )

    doc.add_heading("3. Data Storage and Security", level=1)
    doc.add_paragraph(
        "Personal data is stored in encrypted databases hosted on AWS (us-east-1 region). "
        "Encryption at rest uses AES-256. Data in transit is protected via TLS 1.2 or higher. "
        "Access to production databases is restricted to authorized personnel with role-based "
        "access controls (RBAC)."
    )
    doc.add_paragraph(
        "Backups are performed daily and retained for 90 days. Backup data is encrypted "
        "using the same standards as production data."
    )

    doc.add_heading("4. Data Sharing", level=1)
    doc.add_paragraph(
        "Acme Corp shares personal data with the following categories of third parties: "
        "payment processors (Stripe), cloud infrastructure providers (AWS), analytics "
        "services (Google Analytics), and customer support tools (Zendesk). All third-party "
        "processors are required to sign Data Processing Agreements (DPAs)."
    )

    doc.add_heading("5. Data Retention", level=1)
    doc.add_paragraph(
        "Customer account data is retained for the duration of the business relationship "
        "plus 3 years after account closure. Financial records are retained for 7 years "
        "to comply with tax regulations. Marketing data is retained until consent is withdrawn."
    )

    doc.add_heading("6. User Rights", level=1)
    doc.add_paragraph(
        "Users may request access to their personal data by emailing privacy@acmecorp.com. "
        "We aim to respond to access requests within 30 business days."
    )
    # NOTE: Intentionally missing:
    # - Right to erasure / deletion process
    # - Right to data portability
    # - Right to restrict processing
    # - Right to object to automated decision-making
    # - Process for withdrawing consent

    doc.add_heading("7. Breach Notification", level=1)
    doc.add_paragraph(
        "In the event of a data breach affecting personal data, Acme Corp will notify "
        "affected individuals via email within a reasonable timeframe."
    )
    # NOTE: Intentionally vague — no 72-hour supervisory authority notification,
    # no mention of DPA notification requirements

    doc.add_heading("8. International Transfers", level=1)
    doc.add_paragraph(
        "Acme Corp may transfer personal data to servers located in the United States. "
        "We rely on standard contractual clauses for international data transfers."
    )
    # NOTE: Missing adequacy decision references, no mention of supplementary measures

    doc.add_heading("9. Contact", level=1)
    doc.add_paragraph(
        "For privacy inquiries, contact our privacy team at privacy@acmecorp.com."
    )
    # NOTE: No Data Protection Officer (DPO) designated — required under GDPR Art 37

    out = Path(output_path)
    doc.save(str(out))
    return out


if __name__ == "__main__":
    path = create_sample_policy()
    print(f"Created: {path}")
