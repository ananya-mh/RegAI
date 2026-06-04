from __future__ import annotations

import io
import logging
from typing import Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

RISK_COLORS = {
    "critical": RGBColor(220, 38, 38),
    "high": RGBColor(234, 88, 12),
    "medium": RGBColor(202, 138, 4),
    "low": RGBColor(22, 163, 74),
}

STATUS_COLORS = {
    "compliant": RGBColor(22, 163, 74),
    "partial": RGBColor(202, 138, 4),
    "non-compliant": RGBColor(220, 38, 38),
}


def render_docx(report_sections: list[dict[str, Any]], remediation_tasks: list[dict[str, Any]]) -> bytes:
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    doc.add_heading("Compliance Assessment Report", level=0)

    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer.add_run(
        "AI-generated assessment for informational purposes only. "
        "Does not constitute legal advice."
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    for section in report_sections:
        doc.add_heading(section.get("title", "Compliance Report"), level=1)

        risk = section.get("overall_risk_level", "unknown")
        risk_para = doc.add_paragraph()
        risk_para.add_run("Overall Risk Level: ").bold = True
        risk_run = risk_para.add_run(risk.upper())
        risk_run.bold = True
        risk_run.font.color.rgb = RISK_COLORS.get(risk, RGBColor(0, 0, 0))

        summary = section.get("executive_summary", "")
        if summary:
            doc.add_paragraph(summary)

        findings = section.get("findings", [])
        if findings:
            doc.add_heading("Findings", level=2)
            for finding in findings:
                req = finding.get("requirement", "N/A")
                status = finding.get("status", "unknown")

                finding_heading = doc.add_heading(level=3)
                finding_heading.add_run(f"{req} — ")
                status_run = finding_heading.add_run(status.upper())
                status_run.font.color.rgb = STATUS_COLORS.get(status, RGBColor(0, 0, 0))

                explanation = finding.get("explanation", "")
                if explanation:
                    doc.add_paragraph(explanation)

                recommendation = finding.get("recommendation", "")
                if recommendation:
                    rec_para = doc.add_paragraph()
                    rec_para.add_run("Recommendation: ").bold = True
                    rec_para.add_run(recommendation)

        roadmap = section.get("remediation_roadmap", "")
        if roadmap:
            doc.add_heading("Remediation Roadmap", level=2)
            doc.add_paragraph(roadmap)

    if remediation_tasks:
        doc.add_heading("Remediation Tasks", level=1)

        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        headers = table.rows[0].cells
        headers[0].text = "Priority"
        headers[1].text = "Task"
        headers[2].text = "Description"
        headers[3].text = "Effort"

        for task in remediation_tasks:
            row = table.add_row().cells
            row[0].text = task.get("priority", "medium").upper()
            row[1].text = task.get("title", "N/A")
            row[2].text = (task.get("description", "") or "")[:200]
            row[3].text = task.get("effort_estimate", "TBD") or "TBD"

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def render_pdf(report_sections: list[dict[str, Any]], remediation_tasks: list[dict[str, Any]]) -> bytes:
    html_parts = [
        "<html><head><style>",
        "body { font-family: Calibri, sans-serif; font-size: 11pt; margin: 40px; color: #1a1a1a; }",
        "h1 { color: #1e293b; border-bottom: 2px solid #3b82f6; padding-bottom: 8px; }",
        "h2 { color: #334155; margin-top: 24px; }",
        "h3 { color: #475569; }",
        ".risk-critical { color: #dc2626; font-weight: bold; }",
        ".risk-high { color: #ea580c; font-weight: bold; }",
        ".risk-medium { color: #ca8a04; font-weight: bold; }",
        ".risk-low { color: #16a34a; font-weight: bold; }",
        ".status-compliant { color: #16a34a; }",
        ".status-partial { color: #ca8a04; }",
        ".status-non-compliant { color: #dc2626; }",
        "table { width: 100%; border-collapse: collapse; margin: 16px 0; }",
        "th, td { border: 1px solid #e2e8f0; padding: 8px 12px; text-align: left; font-size: 10pt; }",
        "th { background: #f1f5f9; font-weight: bold; }",
        ".disclaimer { text-align: center; color: #94a3b8; font-size: 9pt; font-style: italic; margin: 20px 0; }",
        "</style></head><body>",
        "<h1>Compliance Assessment Report</h1>",
        '<p class="disclaimer">AI-generated assessment for informational purposes only. Does not constitute legal advice.</p>',
    ]

    for section in report_sections:
        title = section.get("title", "Compliance Report")
        risk = section.get("overall_risk_level", "unknown")
        summary = section.get("executive_summary", "")

        html_parts.append(f"<h2>{title}</h2>")
        html_parts.append(f'<p><strong>Risk Level:</strong> <span class="risk-{risk}">{risk.upper()}</span></p>')
        if summary:
            html_parts.append(f"<p>{summary}</p>")

        findings = section.get("findings", [])
        if findings:
            html_parts.append("<h3>Findings</h3>")
            for f in findings:
                status = f.get("status", "unknown")
                status_class = f"status-{status}"
                html_parts.append(
                    f'<h4>{f.get("requirement", "N/A")} — '
                    f'<span class="{status_class}">{status.upper()}</span></h4>'
                )
                html_parts.append(f'<p>{f.get("explanation", "")}</p>')
                rec = f.get("recommendation", "")
                if rec:
                    html_parts.append(f"<p><strong>Recommendation:</strong> {rec}</p>")

        roadmap = section.get("remediation_roadmap", "")
        if roadmap:
            html_parts.append(f"<h3>Remediation Roadmap</h3><p>{roadmap}</p>")

    if remediation_tasks:
        html_parts.append("<h2>Remediation Tasks</h2>")
        html_parts.append("<table><tr><th>Priority</th><th>Task</th><th>Description</th><th>Effort</th></tr>")
        for task in remediation_tasks:
            html_parts.append(
                f"<tr><td>{task.get('priority', 'medium').upper()}</td>"
                f"<td>{task.get('title', 'N/A')}</td>"
                f"<td>{(task.get('description', '') or '')[:200]}</td>"
                f"<td>{task.get('effort_estimate', 'TBD') or 'TBD'}</td></tr>"
            )
        html_parts.append("</table>")

    html_parts.append("</body></html>")
    html = "\n".join(html_parts)

    try:
        from weasyprint import HTML
        return HTML(string=html).write_pdf()
    except ImportError:
        logger.warning("WeasyPrint not installed, returning HTML as fallback")
        return html.encode("utf-8")
