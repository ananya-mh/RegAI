from __future__ import annotations

import logging
import tempfile
import uuid
from datetime import datetime, timezone
from pathlib import Path

import fitz
from docx import Document as DocxDocument
from fastapi import APIRouter, Depends, HTTPException, UploadFile
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from backend.api.models import PolicyOut, PolicyUploadOut
from backend.models.tables import Policy
from backend.services.database import get_db

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/policies", tags=["policies"])

ALLOWED_EXTENSIONS = {".pdf", ".docx"}
MAX_FILE_SIZE = 50 * 1024 * 1024


def _extract_text_pdf(path: Path) -> tuple[str, list[tuple[int, str]]]:
    doc = fitz.open(str(path))
    full_text = ""
    headers: list[tuple[int, str]] = []
    for page in doc:
        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            if "lines" not in block:
                continue
            for line in block["lines"]:
                spans = line["spans"]
                text = "".join(s["text"] for s in spans)
                if not text.strip():
                    continue
                avg_size = sum(s["size"] for s in spans) / len(spans)
                is_bold = any(s["flags"] & 2 ** 4 for s in spans)
                if avg_size > 13 or is_bold:
                    headers.append((len(full_text), text.strip()))
                full_text += text + "\n"
    doc.close()
    return full_text, headers


def _extract_text_docx(path: Path) -> tuple[str, list[tuple[int, str]]]:
    doc = DocxDocument(str(path))
    full_text = ""
    headers: list[tuple[int, str]] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            headers.append((len(full_text), text))
        full_text += text + "\n"
    return full_text, headers


@router.post("/upload", response_model=PolicyUploadOut)
async def upload_policy(
    file: UploadFile,
    db: AsyncSession = Depends(get_db),
) -> PolicyUploadOut:
    if not file.filename:
        raise HTTPException(400, "No filename provided")

    ext = Path(file.filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(400, f"Unsupported file type: {ext}. Use PDF or DOCX.")

    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(400, f"File too large (max {MAX_FILE_SIZE // 1024 // 1024}MB)")

    with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
        tmp.write(content)
        tmp_path = Path(tmp.name)

    try:
        if ext == ".pdf":
            text, headers = _extract_text_pdf(tmp_path)
        else:
            text, headers = _extract_text_docx(tmp_path)
    finally:
        tmp_path.unlink(missing_ok=True)

    if not text.strip():
        raise HTTPException(400, "Could not extract any text from the file")

    parsed_path = f"policies/{uuid.uuid4()}.txt"
    parsed_dir = Path("data") / "policies"
    parsed_dir.mkdir(parents=True, exist_ok=True)
    (Path("data") / parsed_path).write_text(text, encoding="utf-8")

    policy = Policy(
        id=uuid.uuid4(),
        filename=file.filename,
        upload_date=datetime.now(timezone.utc),
        parsed_text_path=parsed_path,
    )
    db.add(policy)
    await db.flush()

    try:
        from backend.rag.chunking import chunk_policy
        from backend.rag.vector_stores import PolicyStore
        from backend.services.config import settings

        upload_date = policy.upload_date.isoformat() if policy.upload_date else ""
        chunks = chunk_policy(
            text=text,
            document_name=file.filename,
            upload_date=upload_date,
            section_headers=headers,
        )
        store = PolicyStore(persist_path=settings.chroma_persist_path)
        store.add(chunks)
    except Exception:
        logger.warning("Failed to chunk/index in Chroma", exc_info=True)
        chunks = []

    await db.commit()
    await db.refresh(policy)

    return PolicyUploadOut(
        policy=PolicyOut.model_validate(policy),
        chunks_created=len(chunks),
    )


@router.get("", response_model=list[PolicyOut])
async def list_policies(
    db: AsyncSession = Depends(get_db),
) -> list[PolicyOut]:
    result = await db.execute(select(Policy).order_by(Policy.upload_date.desc()))
    policies = result.scalars().all()
    return [PolicyOut.model_validate(p) for p in policies]
